"""Extract plain text from an uploaded proposal file.

Used by the EiR reviewer so a PI can upload their actual proposal files instead
of pasting. Deliberately small and dependency-free beyond what the app already
pins (pdfplumber, python-docx — both already in requirements.txt).

TWO RULES, both learned the hard way in this repo:

1. **NEVER TRUNCATE SILENTLY.** CLAUDE.md records `MAX_PDF_PAGES = 40` /
   `MAX_CHARS = 200_000` quietly returning 36% of a 118-page policy document as
   clean text with no error, so the document "looked complete" and a coverage
   check wrongly cleared. Here the consequence would be worse than a bad audit:
   a truncated read makes the LAST sections of a proposal look MISSING, and the
   PI would go rewrite a sustainability plan they already wrote. So the caps
   below are generous, and hitting one sets a `truncated` flag that the caller
   MUST surface. A truncated read is reported, never swallowed.

2. **AN UNREADABLE FILE IS NOT AN EMPTY FILE.** A failed parse returns an
   `error` string, not "". The reviewer reports it as a file it could not read,
   never as content the author omitted — same invariant as the KB scraper's
   `looks_unreadable()`.
"""

from __future__ import annotations

import io
import os
import re
from typing import Optional

# Generous, and enforced VISIBLY. A 15-page NSF Project Description is ~10k
# words; a full package with letters and budget justification is well under
# these. Anything past them is almost certainly the wrong file.
MAX_PAGES = 400
MAX_CHARS = 1_500_000

# How pdfplumber decides where one word ends and the next begins.
#
# It measures the horizontal GAP between glyphs against a tolerance, and that
# tolerance defaults to a FIXED 3 points. TeX does not write a space character
# between words — it moves the pen — and in a 10pt document that move is often
# ~2pt, under the default. The words are then welded together.
#
# Measured on a real awarded NSF 23-598 package (Morgan State, 56 pages,
# LaTeX): 863 run-together tokens over 18 characters, including every heading
# the reviewer looks for — `IntellectualMerit`, `BroaderImpacts`,
# `ResultsfromPriorNSFsupport`. Nothing reported a broken read; three separate
# features just said false things about the PROPOSAL instead. The locate stage
# could not find those sections; `quote_in` rejected the reviewer's own quotes,
# so 21 of 34 fix-list rows were correct findings demoted to `not_found`; and
# the language checks returned 142 "mistakes" of which exactly one was real.
#
# A RATIO scales the tolerance with the font size instead of pinning it, which
# is the property that makes it safe across documents: 0.15 is 1.5pt at 10pt
# body text and proportionally more in a heading. Measured across both awarded
# packages: run-together tokens 863 -> 1 and 21 -> 1, all heading probes
# recovered, and the whole-word counts of ordinary vocabulary UNCHANGED
# (representation 85, mathematics 78, undergraduate 60, research 269) — so the
# spacing was not bought by chopping real words up, which is the mirror failure
# and would break `quote_in` just as badly while looking like a different bug.
#
# 0.20 already regressed (111 welded tokens); 0.25 regressed badly (613).
# Shared with services/solicitation_extractor.read_pdf so the upload path and
# the solicitation path cannot drift into reading the same PDF differently.
PDF_X_TOLERANCE_RATIO = 0.15

# THE VERTICAL HALF OF THE SAME PROBLEM, and it scrambles READING ORDER rather
# than spacing. pdfplumber groups characters into lines by baseline, and
# pdfplumber's default y_tolerance of 3 splits a line whenever part of it sits
# slightly off — a superscript-styled numeral, or an italic run set in a
# different face. Measured on the awarded NSF EiR Project Description, where
# both shapes occur:
#
#   "...two undergraduates in Year ; three \n 1 \n undergraduates in Year ; and,"
#   "a part of my \n Super Representation \n with two primary goals: ... \n
#    Theory research program \n representation theory of..."
#
# The digits and the italic phrase are lifted OUT of their sentences onto lines
# of their own. The reviewer reassembles the sentence correctly and quotes it,
# `quote_in` compares against our scrambled copy, fails, and golden rule 2
# demotes a real `addressed` to `not_found` — so a funded proposal was told it
# never described its undergraduate research opportunities, ten runs out of ten.
#
# 5 not 8: both fix the two cases above, but 8 also MERGES lines that are
# genuinely separate (+1,117 chars of run-together text on the same document)
# while 5 changes the length by 0.1% (57,763 -> 57,682). Shared with
# solicitation_extractor.read_pdf for the same reason the ratio above is: the
# upload path and the solicitation path must not read one PDF differently.
PDF_Y_TOLERANCE = 5

# Extensions we can actually read. Checked against the filename only as a first
# pass — content sniffing decides for anything ambiguous.
PDF_EXTS = {".pdf"}
DOCX_EXTS = {".docx"}
TEXT_EXTS = {".txt", ".md", ".markdown", ".text", ".rst", ".csv", ".tex"}
SUPPORTED_EXTS = PDF_EXTS | DOCX_EXTS | TEXT_EXTS

# Formats that need a dependency this app does not pin. Named explicitly so the
# PI gets "convert this to PDF" instead of a generic failure.
KNOWN_UNSUPPORTED = {
    ".doc": "legacy Word (.doc)",
    ".rtf": "rich text (.rtf)",
    ".pages": "Apple Pages",
    ".odt": "OpenDocument (.odt)",
    ".wpd": "WordPerfect",
}


def _extract_pdf(data: bytes):
    """(text, page_count, truncated, page_texts). Raises on an unparseable PDF.

    `page_texts` is returned rather than discarded so `services.pdf_sections`
    can map page ranges onto the very string built here — the two must not
    extract the document twice and disagree about it."""
    import pdfplumber

    pages: list[str] = []
    truncated = False
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        total = len(pdf.pages)
        for i, page in enumerate(pdf.pages):
            if i >= MAX_PAGES:
                truncated = True
                break
            pages.append(
                page.extract_text(x_tolerance_ratio=PDF_X_TOLERANCE_RATIO,
                                  y_tolerance=PDF_Y_TOLERANCE) or "")
    text = "\n".join(pages)
    if len(text) > MAX_CHARS:
        text = text[:MAX_CHARS]
        truncated = True
    return text, total, truncated, pages


def _extract_docx(data: bytes) -> tuple[str, int, bool]:
    """(text, page_count, truncated). Page count is 0 — a .docx has no fixed
    pagination until it is rendered, and inventing one would feed a bogus number
    into the 2-page letter check."""
    import docx

    doc = docx.Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    # Tables carry real content in budget justifications and timelines; skipping
    # them would drop exactly the numbers a reviewer is looking for.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    text = "\n".join(parts)
    truncated = len(text) > MAX_CHARS
    return (text[:MAX_CHARS] if truncated else text), 0, truncated


def _extract_plain(data: bytes) -> tuple[str, int, bool]:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            text = data.decode(encoding)
            break
        except (UnicodeDecodeError, LookupError):
            continue
    else:
        text = data.decode("utf-8", errors="replace")
    truncated = len(text) > MAX_CHARS
    return (text[:MAX_CHARS] if truncated else text), 0, truncated


def _looks_like_pdf(data: bytes) -> bool:
    return data[:5] == b"%PDF-"


def _looks_like_zip(data: bytes) -> bool:
    # .docx is a zip. So is .pptx/.xlsx, which is why the caller still checks
    # the extension before trusting this.
    return data[:2] == b"PK"


def extract_upload(filename: str, data: bytes, *, sections: dict = None) -> dict:
    """Read one uploaded file.

    Returns {filename, text, pages, chars, truncated, error}. NEVER raises —
    an unreadable file comes back with `error` set and `text` empty, so one bad
    file in a multi-file upload cannot take down the whole review.

    `sections` opts one file into STRUCTURAL splitting: when a PDF turns out to
    be an assembled multi-attachment package (a Research.gov submission), the
    result also carries `section_spans`, offsets into this file's own `text`.
    Absent, unsplittable, or unsafe to split, the key is simply not there and
    nothing downstream changes — see services/pdf_sections for the fail-safes."""
    name = filename or "file"
    ext = os.path.splitext(name)[1].lower()
    out = {"filename": name, "text": "", "pages": 0, "chars": 0,
           "truncated": False, "error": None, "page_texts": []}

    if not data:
        out["error"] = "The file is empty."
        return out

    if ext in KNOWN_UNSUPPORTED:
        out["error"] = (f"{KNOWN_UNSUPPORTED[ext]} isn't supported. "
                        "Save it as PDF or .docx and try again.")
        return out

    try:
        # Sniff content first: a mislabelled extension is common when files come
        # off a shared drive, and the bytes are authoritative.
        if _looks_like_pdf(data):
            text, pages, truncated, page_texts = _extract_pdf(data)
        elif ext in DOCX_EXTS and _looks_like_zip(data):
            text, pages, truncated = _extract_docx(data)
            page_texts = []
        elif ext in TEXT_EXTS or not ext:
            text, pages, truncated = _extract_plain(data)
            page_texts = []
        elif ext in PDF_EXTS:
            out["error"] = "This is named .pdf but isn't a PDF file."
            return out
        elif _looks_like_zip(data):
            out["error"] = ("This looks like an Office file that isn't .docx "
                            "(maybe .pptx or .xlsx). Save it as PDF and try again.")
            return out
        else:
            out["error"] = (f"Can't read {ext or 'this file type'}. "
                            "Upload a PDF, .docx, or plain text file.")
            return out
    except Exception as e:                      # noqa: BLE001 — never raise to the caller
        out["error"] = f"Couldn't read this file ({type(e).__name__})."
        return out

    text = text.strip()
    if not text:
        # A real failure mode: a scanned/photographed proposal is a PDF of
        # IMAGES with no text layer. Say so, because "0 words" alone reads as an
        # empty file and the PI will re-upload the same thing.
        out["error"] = ("No text found. If this is a scanned PDF, it has no text "
                        "layer — export it from the original document instead.")
        return out

    out.update(text=text, pages=pages, chars=len(text), truncated=truncated)

    # Returned rather than dropped: `services.page_ledger` accounts for every
    # page against this very list, and re-extracting the PDF to get it back
    # would risk the two reads disagreeing -- the failure `_extract_pdf`'s
    # docstring warns about.
    out["page_texts"] = page_texts or []

    # STRUCTURAL SPLIT, opt-in and fail-safe. Only here, where the raw bytes and
    # the per-page texts are both still in scope, so nothing has to carry a
    # 60 MB PDF through the mapping layer.
    if sections and page_texts and not truncated:
        try:
            from services import pdf_sections as _ps
            from services import page_ledger as _pl

            spans, report = _ps.split(data, page_texts, sections)
            shift = len(_raw) - len(_raw.lstrip()) if (_raw := "\n".join(page_texts)) else 0

            # STRUCTURE FIRST, and it wins. `pdf_sections` reads the seams out
            # of the PDF's object graph and returns the same answer every run;
            # the walk fills the pages it could not name and never overrules it.
            structure = {}
            for key, span in (spans or {}).items():
                for page, (p0, p1) in enumerate(_pl._page_offsets(page_texts), start=1):
                    if span["start"] <= p0 and p1 <= span["end"]:
                        structure[page] = key

            ledger = _pl.build_ledger(page_texts, sections, structure=structure)
            out["page_ledger"] = ledger
            out["ledger_toc_mismatch"] = _pl.reconcile_toc(ledger, page_texts, sections)

            merged = _pl.spans_from_ledger(ledger, page_texts, sections)
            out["ledger_page_counts"] = _pl.page_counts_from_ledger(ledger)

            # `text` was stripped after joining, so every offset shifts by
            # the leading whitespace that removal took out. Clamped, and the
            # slice is re-read from the FINAL text so a caller can never be
            # handed an offset that does not address what it claims to.
            rebased = {}
            for key, span in merged.items():
                s0 = max(0, span["start"] - shift)
                e0 = max(s0, min(len(text), span["end"] - shift))
                if e0 > s0:
                    rebased[key] = {**span, "start": s0, "end": e0,
                                    "text": text[s0:e0]}
            if rebased:
                out["section_spans"] = rebased
                out["section_report"] = report
        except Exception as exc:                # never break an upload over this
            print(f"[DOCUMENT-TEXT] structural split skipped: {exc}")
    return out


# A NUMBERED PREFIX IS ORDERING, NOT A NAME. Real packages arrive as
# "01-Project-Summary.pdf" ... "11-Letters-and-Supplementary-Documents.pdf".
_FILE_ORDER_PREFIX = re.compile(r"^\s*\d+[-_. ]+")


def _section_from_filename(sections: dict, heading: str):
    """(section key, how) for a FILENAME read as a section name, or (None, None).

    `how` is "filename" | "filename_subset" | "filename_narrowed".

    WHY THIS IS NOT `resolve_section_key`. That matcher requires set EQUALITY of
    meaning-carrying words, and it must keep requiring it: it is what stops
    "Project Description Supplementary Documents" folding into
    `project_description` and losing a real section, guarded by three tests in
    three files. But equality is too strict for a FILENAME, which a PI writes
    for themselves. Measured on a real awarded package, only 5 of 11 filenames
    resolved, and the miss cost the score's denominator: the four
    `supplementary_document` rules of `11-Letters-and-Supplementary-Documents.pdf`
    came back `could_not_locate` in 4 of 5 identical runs, moving `assessed`
    45 <-> 49, because {letter, supplementary, document} != {supplementary,
    document}.

    So the widening lives HERE, where the only input is a filename, rather than
    in the shared matcher the locate stage and `generic_checks` also call. Same
    move as `rulebook_checks`, which strips markdown from its own probe rather
    than widening the shared `heading_regex`.

    Three tiers, first hit wins, each refusing when two sections qualify —
    ambiguity falls through to the locate stage, which is today's behaviour.
    """
    # Imported inside, like `map_files_to_sections` does, so this module stays
    # importable on its own.
    from services import solicitation_profile as _sp

    if not sections or not heading:
        return None, None

    exact = _sp.resolve_section_key(sections, heading)
    if exact:
        return exact, "filename"

    file_sig = _sp.section_signature(heading)
    if not file_sig:
        return None, None

    # TIER 2 — the filename says MORE than the section is called.
    #
    # THE COVERAGE GUARD IS WHAT MAKES THIS SOUND, and "exactly one subset
    # match" is NOT sound without it. That rule's safety in the
    # Project-Description case comes from the universe happening to hold BOTH
    # sections; a universe holding a section signed {plan} would swallow
    # "Data Management and Sharing Plan" outright, with no ambiguity to refuse
    # on. Requiring the matched name to account for at least HALF the filename's
    # meaning-words bounds it: a match can only claim a file it genuinely shares
    # most of its name with. It costs the odd true positive (a one-word section
    # named by a long filename), and losing one falls back to today's behaviour,
    # which is the safe direction.
    hits = [k for k, meta in sections.items()
            if any(sig < file_sig and 2 * len(sig) >= len(file_sig)
                   for sig in _sp.section_signatures(k, meta))]
    if len(hits) == 1:
        return hits[0], "filename_subset"
    if hits:
        return None, None            # ambiguous: two sections both fit

    # TIER 3 — the filename says LESS ("Mentoring Plan" for a section called
    # "Postdoctoral Mentoring Plan").
    #
    # THE ONLY TIER HERE THAT CAN PRODUCE A CONFIDENT WRONG VERDICT. Every other
    # match can at worst move a rule from unassessed to assessed against text
    # that really is that section; this one can file a GRADUATE mentoring plan
    # under a POSTDOCTORAL mentoring plan section, where its rules are then
    # judged. Hence: at most one missing word, never from a single-word
    # filename, exactly one candidate — and it reports itself as
    # `filename_narrowed` rather than `filename`, so the guess is visible in the
    # extraction report and one dropdown from the `chosen` override.
    if len(file_sig) < 2:
        return None, None
    near = [k for k, meta in sections.items()
            if any(file_sig < sig and len(sig - file_sig) == 1
                   for sig in _sp.section_signatures(k, meta))]
    if len(near) == 1:
        return near[0], "filename_narrowed"
    return None, None


def _stem_as_section_name(filename: str) -> str:
    """The filename read as a section NAME: ordering stripped, separators spaced."""
    stem = os.path.splitext(filename or "")[0]
    stem = _FILE_ORDER_PREFIX.sub("", stem)
    return stem.replace("_", " ").replace("-", " ").strip()


def map_files_to_sections(files: list[dict], sections: dict) -> tuple[str, dict, list, list]:
    """Which uploaded file IS which section, so the reviewer need not guess.

    Returns `(text, spans, leftover, mapping)`:
      text     — every file joined, filename as a heading, exactly as `combine()`
      spans    — {section key: span} with start/end as REAL offsets into `text`
      leftover — files that mapped to nothing, still inside `text` for the locate stage
      mapping  — one row per file, for the UI, so a mis-map is visible not silent

    WHY. Measured over five uploads of one awarded 11-file package: `locate_sections`
    found 6 sections on one run and ONE on another -- reading the entire 45-page
    package as "References Cited", which collapsed 48 assessable rules to 14 and
    scored a FUNDED proposal at 29%. The seams were never in doubt: the PI uploaded
    one file per section. This hands them over instead of paying a model call to
    rediscover them.

    A HYBRID, NEVER A REPLACEMENT. Only ~5 of 11 real filenames resolve, so anything
    unmapped stays in `text` and reaches the reviewer through `locate_sections` as
    before -- the same path a pasted draft and a single combined PDF depend on.

    OFFSETS ARE INTO THE COMBINED TEXT, not into the file. `draft_review` slices by
    `span["start"]`/`["end"]` to carve Broader Impacts out of the Project Description
    and to order the section map, so a span carrying file-local offsets would corrupt
    both. Building the text here is what lets the offsets be right.

    Resolution reuses `solicitation_profile.resolve_section_key`, so a filename is
    matched by the SAME rule a requirement row is, and the matcher is never loosened
    here (its set-equality is what stops "Project Description Supplementary
    Documents" folding into "Project Description").
    """
    from services import solicitation_profile as _sp

    spans: dict = {}
    leftover: list = []
    mapping: list = []
    chunks: list = []
    cursor = 0
    for f in files or []:
        row = {"filename": f.get("filename"), "section": None, "source": None,
               "pages": f.get("pages") or 0}
        text = (f.get("text") or "").strip()
        if f.get("error") or not text:
            # No text is no section. Already reported in `extraction.files`.
            row["source"] = "unreadable"
            mapping.append(row)
            continue

        heading = _stem_as_section_name(f.get("filename") or "")
        chunk = f"{heading}\n\n{text}" if heading else text
        chunk_full = chunk
        # Where this file's own text begins once the heading and the joiner are in.
        body_at = cursor + (len(heading) + 2 if heading else 0)

        chosen = f.get("section")
        key = None
        if chosen and chosen in (sections or {}):
            key, row["source"] = chosen, "chosen"
        else:
            key, how = _section_from_filename(sections or {}, heading)
            if key:
                row["source"] = how

        # ONE FILE, MANY SECTIONS — the assembled Research.gov package. When the
        # filename named nothing but the PDF's own structure did, file every
        # section it identified. Same first-claim-wins rule, so a loose section
        # file uploaded alongside still keeps whatever it claimed first.
        if key is None and f.get("section_spans"):
            placed = []
            for sec_key, sec in f["section_spans"].items():
                if sec_key in spans or sec_key not in (sections or {}):
                    continue
                s0, e0 = body_at + sec["start"], body_at + sec["end"]
                spans[sec_key] = {
                    # Sliced from the combined text, never from the file's copy,
                    # so `text[start:end] == span["text"]` holds by construction
                    # rather than by two pieces of arithmetic agreeing.
                    "text": chunk_full[s0 - cursor:e0 - cursor],
                    "start": s0, "end": e0,
                    "marker": sec.get("marker") or sec_key,
                    "pages": sec.get("pages"), "filename": f.get("filename"),
                }
                placed.append(sec_key)
            if placed:
                row["source"] = "pdf_structure"
                row["sections"] = placed          # one file is not one section here

        # A section already claimed keeps its first file. Two files for one section
        # is a packaging mistake worth surfacing; silently overwriting one is the
        # failure mode this codebase keeps having to undo.
        if key and key not in spans:
            spans[key] = {
                "text": text, "start": body_at, "end": body_at + len(text),
                "marker": heading or key,
                "pages": f.get("pages") or 0, "filename": f.get("filename"),
            }
            row["section"] = key
        else:
            if key:
                row["source"] = "duplicate"
            leftover.append(f)

        mapping.append(row)
        chunks.append(chunk)
        cursor += len(chunk) + 2      # the "\n\n" joiner below

    return "\n\n".join(chunks), spans, leftover, mapping


def combine(files: list[dict]) -> str:
    """Join several extracted files into one document for the reviewer.

    Each file's name is emitted as a heading line. That is not cosmetic: the EiR
    reviewer segments by heading, and a file called "Letter of Institutional
    Support.pdf" gives the locate stage a perfect marker for a section that is
    otherwise easy to miss. The stem is used verbatim so it reads as a heading
    rather than a path."""
    chunks = []
    for f in files:
        if f.get("error") or not f.get("text"):
            continue
        stem = os.path.splitext(f["filename"])[0].replace("_", " ").strip()
        chunks.append(f"{stem}\n\n{f['text']}")
    return "\n\n".join(chunks)
