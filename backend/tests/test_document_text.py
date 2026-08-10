"""Tests for services/document_text.py — extracting text from uploaded files.

The invariant under test throughout: an unreadable file must be reported as
unreadable, never as an empty or absent document.
"""

import io

import pytest

from services import document_text as dt


def _pdf_bytes(pages: list[str]) -> bytes:
    """A real, minimal PDF with a text layer, built with pypdf + reportlab if
    available; otherwise skip (we refuse to assert against a fake)."""
    reportlab = pytest.importorskip("reportlab", reason="reportlab not installed")
    from reportlab.pdfgen import canvas            # noqa: PLC0415
    from reportlab.lib.pagesizes import letter     # noqa: PLC0415

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    for text in pages:
        c.drawString(72, 720, text)
        c.showPage()
    c.save()
    return buf.getvalue()


def _docx_bytes(paragraphs: list[str], table: list[list[str]] | None = None) -> bytes:
    import docx

    buf = io.BytesIO()
    d = docx.Document()
    for p in paragraphs:
        d.add_paragraph(p)
    if table:
        t = d.add_table(rows=len(table), cols=len(table[0]))
        for r, row in enumerate(table):
            for col, val in enumerate(row):
                t.cell(r, col).text = val
    d.save(buf)
    return buf.getvalue()


# ── plain text ──────────────────────────────────────────────────────────────

def test_plain_text_is_read():
    out = dt.extract_upload("draft.txt", b"Project Description\nOur goal is X.")
    assert out["error"] is None
    assert "Our goal is X." in out["text"]
    assert out["chars"] > 0


def test_markdown_is_read():
    out = dt.extract_upload("draft.md", b"# Project Description\n\nOur goal.")
    assert out["error"] is None
    assert "Our goal." in out["text"]


def test_utf16_text_is_decoded():
    out = dt.extract_upload("draft.txt", "Sustainability plan".encode("utf-16"))
    assert out["error"] is None
    assert "Sustainability plan" in out["text"]


def test_undecodable_bytes_do_not_raise():
    out = dt.extract_upload("draft.txt", b"\xff\xfe\x00 broken \xc3\x28")
    assert out["error"] is None or out["text"] == ""


# ── docx ────────────────────────────────────────────────────────────────────

def test_docx_paragraphs_are_read():
    data = _docx_bytes(["Project Description", "Our overall research goal is X."])
    out = dt.extract_upload("draft.docx", data)
    assert out["error"] is None
    assert "Our overall research goal is X." in out["text"]


def test_docx_tables_are_read():
    """Budget justifications and timelines put the real numbers in tables;
    skipping them would drop exactly what a reviewer looks for."""
    data = _docx_bytes(["Budget Justification"],
                       table=[["Category", "Amount"], ["Equipment", "$58,000"]])
    out = dt.extract_upload("budget.docx", data)
    assert "$58,000" in out["text"]


def test_docx_reports_no_page_count():
    """A .docx has no fixed pagination until rendered. Inventing one would feed
    a bogus number into the 2-page institutional letter check."""
    out = dt.extract_upload("draft.docx", _docx_bytes(["hello"]))
    assert out["pages"] == 0


# ── pdf ─────────────────────────────────────────────────────────────────────

def test_pdf_text_and_page_count():
    data = _pdf_bytes(["Project Description", "Sustainability plan", "References"])
    out = dt.extract_upload("draft.pdf", data)
    assert out["error"] is None
    assert out["pages"] == 3
    assert "Sustainability plan" in out["text"]


def test_pdf_is_detected_by_content_not_extension():
    """Files off a shared drive are routinely mislabelled; the bytes win."""
    data = _pdf_bytes(["Project Description"])
    out = dt.extract_upload("proposal.txt", data)      # wrong extension
    assert out["error"] is None
    assert "Project Description" in out["text"]


def test_file_named_pdf_that_is_not_a_pdf_is_an_error_not_empty():
    out = dt.extract_upload("draft.pdf", b"this is plainly not a pdf at all")
    assert out["error"]
    assert out["text"] == ""


# ── unreadable is not empty ─────────────────────────────────────────────────

def test_empty_file_is_an_error():
    out = dt.extract_upload("draft.pdf", b"")
    assert out["error"] == "The file is empty."


def test_legacy_doc_names_the_format_and_says_what_to_do():
    out = dt.extract_upload("draft.doc", b"\xd0\xcf\x11\xe0legacy ole2")
    assert "legacy Word" in out["error"]
    assert "PDF" in out["error"]


@pytest.mark.parametrize("ext", [".rtf", ".pages", ".odt", ".wpd"])
def test_other_unsupported_formats_are_named(ext):
    out = dt.extract_upload(f"draft{ext}", b"some bytes")
    assert out["error"]
    assert out["text"] == ""


def test_scanned_pdf_with_no_text_layer_says_so():
    """A photographed proposal is a PDF of images. '0 words' alone reads as an
    empty file and the PI re-uploads the same thing."""
    data = _pdf_bytes([""])
    out = dt.extract_upload("scan.pdf", data)
    assert out["error"]
    assert "scanned" in out["error"].lower()


def test_pptx_is_rejected_with_a_useful_message():
    """.pptx is also a zip, so it must not be mistaken for .docx."""
    import zipfile
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as z:
        z.writestr("ppt/presentation.xml", "<p/>")
    out = dt.extract_upload("slides.pptx", buf.getvalue())
    assert out["error"]
    assert "PDF" in out["error"]


# ── truncation is reported, never silent ────────────────────────────────────

def test_oversized_text_sets_the_truncated_flag(monkeypatch):
    """CLAUDE.md: a silent truncation makes the LAST sections look missing, and
    the PI goes and rewrites something they already wrote."""
    monkeypatch.setattr(dt, "MAX_CHARS", 100)
    out = dt.extract_upload("draft.txt", b"x" * 500)
    assert out["truncated"] is True
    assert out["chars"] == 100


def test_normal_sized_file_is_not_flagged_truncated():
    out = dt.extract_upload("draft.txt", b"Project Description\nShort and complete.")
    assert out["truncated"] is False


def test_pdf_page_cap_sets_truncated_and_keeps_the_true_total(monkeypatch):
    monkeypatch.setattr(dt, "MAX_PAGES", 2)
    data = _pdf_bytes(["one", "two", "three", "four"])
    out = dt.extract_upload("long.pdf", data)
    assert out["truncated"] is True
    assert out["pages"] == 4          # the real total, not the number we read
    assert "four" not in out["text"]


# ── combine ─────────────────────────────────────────────────────────────────

def test_combine_uses_filenames_as_headings():
    """Not cosmetic: the EiR reviewer segments by heading, so a file called
    'Letter of Institutional Support.pdf' hands the locate stage a marker for a
    section that is otherwise easy to miss."""
    files = [
        {"filename": "Project Description.pdf", "text": "Our goal.", "error": None},
        {"filename": "Letter of Institutional Support.pdf", "text": "The dean writes.",
         "error": None},
    ]
    combined = dt.combine(files)
    assert "Project Description\n" in combined
    assert "Letter of Institutional Support\n" in combined
    assert "The dean writes." in combined


def test_combine_underscores_become_spaces():
    files = [{"filename": "Budget_Justification.docx", "text": "x", "error": None}]
    assert "Budget Justification" in dt.combine(files)


def test_combine_skips_unreadable_files_without_dropping_the_rest():
    files = [
        {"filename": "good.pdf", "text": "Real content.", "error": None},
        {"filename": "bad.pdf", "text": "", "error": "Couldn't read this file."},
    ]
    combined = dt.combine(files)
    assert "Real content." in combined
    assert "bad" not in combined


def test_combine_of_nothing_readable_is_empty():
    files = [{"filename": "bad.pdf", "text": "", "error": "nope"}]
    assert dt.combine(files).strip() == ""
